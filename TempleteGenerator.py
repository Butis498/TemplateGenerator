import os
import json
import pandas as pd
from docx import Document
import openpyxl
from openpyxl import load_workbook
from pathlib import Path
from xlrd import open_workbook
from xlutils.copy import copy
import xlsxwriter

class TempleteGenerator():

    def __init__(self):
        self.Keydictionary = ["REPLACE_NAME" , "REPLACE_LASTNAME", "REPLACE_CITY", "REPLACE_DIRECTION" , "REPLACE_POSSTALCODE" , "REPLACE_PHONENUM", "REPLACE_SOCIALCODE",
                            "REPLACE_EMAIL" , "REPLACE_COUNTRY" , "REPLACE_CARDNUM" , "REPLACE_EXPIRATION" , "REPLACE_CVV" , "REPLACE_ACTIVES" , "REPLACE_PASIVES" ,
                            "REPLACE_HERITAGE" , "REPLACE_COMPANY" , "REPLACE_DATE" , "REPLACE_ACTIVITY" , "REPLACE_RUC" , "REPLACE_ACCOUNT" , "REPLACE_POSITION",
                            "REPLACE_INSTITUTION" , "REPLACE_PERIOD" , "REPLACE_ABA" , "REPLACE_SWIFT"]

        self.json_data = {"REPLACE_NAME":"NOMBRE" , "REPLACE_LASTNAME":"APELLIDO", "REPLACE_CITY":"CIUDAD", "REPLACE_DIRECTION":"DIRECCION" , "REPLACE_POSSTALCODE":"CODIGO POSTAL", "REPLACE_PHONENUM":"TELEFONO", "REPLACE_SOCIALCODE":"SEGURO SOCIAL",
                            "REPLACE_EMAIL":"MAIL" , "REPLACE_COUNTRY":"PAIS" , "REPLACE_CARDNUM":"CARD NUMBER" , "REPLACE_EXPIRATION":"CADUCIDAD" , "REPLACE_CVV":"CVV" , "REPLACE_ACTIVES":"ACTIVOS" , "REPLACE_PASIVES": "PASIVOS" ,
                            "REPLACE_HERITAGE":"PATRIMONIO" , "REPLACE_COMPANY":"EMPRESA" , "REPLACE_DATE":"FECHA" , "REPLACE_ACTIVITY":"ACTIVIDAD" , "REPLACE_RUC":"RUC" , "REPLACE_ACCOUNT":"CUENTA" , "REPLACE_POSITION": "CARGO",
                            "REPLACE_INSTITUTION": "INSTITUCION" , "REPLACE_PERIOD":"PERIODO" , "REPLACE_ABA":"ABA" , "REPLACE_SWIFT":"SWIFT"}


        self.db =pd.read_csv('DP_Ecuador_10.csv') 
        self.file = open("DP_Ecuador_10.csv")
        self.data_count = len(self.file.readlines()) - 1
        self.file.close()
        self.Generated_folder_name = "\\Generated"
        self.path = str(os.getcwd() + self.Generated_folder_name)


        if not os.path.isdir(os.getcwd() + self.Generated_folder_name):
            os.mkdir(os.getcwd() + self.Generated_folder_name)


    def replacedParograph(self , parograph , key , data_to_replace):
        try:
            replaced = parograph.replace(key , data_to_replace)
        except TypeError as _:
            replaced = ""

        return replaced


    def get_Data(self , i , key):
        return str(self.db[self.json_data[key]].iloc[i])

    def isfloat(self , value):
        try:
            float(value)
            return True
        except ValueError:
            return False

    def replaceContentXls(self ,documentName ):

        for i in range(self.data_count):
            
            document = load_workbook(documentName)

            for sheet in document._sheets:

                for row in range(1,sheet.max_row + 1):
                    for col in range(1,sheet.max_column + 1):
                        par = str(sheet.cell(row,col).value)
                        if par == 'None':
                            par = ""
                            
                        for key in self.Keydictionary:
                            if key in par:
                                data_to_replace = self.get_Data(i , key)
                                par = self.replacedParograph(par , key , data_to_replace )
                            else:
                                if self.isfloat(par):
                                    float_value = float(par)
                                    sheet.cell(row, col).value = float_value
                                else:
                                    sheet.cell(row, col).value = par

                        if self.isfloat(par):
                            float_value = float(par)
                            sheet.cell(row, col).value = float_value
                        else:
                            sheet.cell(row, col).value = par

            path, file = os.path.split(documentName)
            file_name , file_extension = os.path.splitext(file)
            path = Path(path)
            path = str(path.parent)
            document.save(self.path  + "\\"+ file_name + "_" + str(i) + file_extension)
        


    def replaceContentDoc(self , documentName ):

        for i in range(self.data_count):
            document = Document(documentName)

            for paragraph in document.paragraphs:
                for key in self.Keydictionary:
                    if key in paragraph.text:
                        data_to_replace = self.get_Data(i , key)
                        paragraph.text = self.replacedParograph(paragraph.text , key , data_to_replace)
            for table in document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for key in self.Keydictionary:
                                if key in paragraph.text:
                                    data_to_replace = self.get_Data(i , key)
                                    paragraph.text = self.replacedParograph(paragraph.text , key , data_to_replace)
            path, file = os.path.split(documentName)
            file_name , file_extension = os.path.splitext(file)
            path = Path(path)
            path = str(path.parent)
            document.save(self.path  + "\\"+ file_name + "_" + str(i) + file_extension)
                    
    


